from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)

# ── Styles ──────────────────────────────────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

h1 = doc.styles['Heading 1']
h1.font.name = 'Calibri'
h1.font.size = Pt(16)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

h2 = doc.styles['Heading 2']
h2.font.name = 'Calibri'
h2.font.size = Pt(13)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

h3 = doc.styles['Heading 3']
h3.font.name = 'Calibri'
h3.font.size = Pt(11)
h3.font.bold = True
h3.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def placeholder(text):
    """Gray italic screenshot placeholder paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[ SCREENSHOT: {text} ]')
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.size = Pt(10)
    # Light gray shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    doc.add_paragraph()  # spacing after

def note(text):
    p = doc.add_paragraph()
    run = p.add_run('NOTE: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x55, 0x00)
    run2 = p.add_run(text)
    run2.font.color.rgb = RGBColor(0xC0, 0x55, 0x00)

def tip(text):
    p = doc.add_paragraph()
    run = p.add_run('TIP: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x7A, 0x1F)
    run2 = p.add_run(text)
    run2.font.color.rgb = RGBColor(0x1F, 0x7A, 0x1F)

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))

def numbered(text):
    doc.add_paragraph(text, style='List Number')

# ════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('CATCAT')
run.font.name = 'Calibri'
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Computer-Assisted Target Classification and Annotation Tool')
run.font.name = 'Calibri'
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
doc.add_paragraph()

doctype = doc.add_paragraph()
doctype.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = doctype.add_run('STANDARD OPERATING PROCEDURE')
run.font.name = 'Calibri'
run.font.size = Pt(18)
run.font.bold = True

doc.add_paragraph()

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver.add_run('Training & Operations Guide\nVersion 1.0  |  June 2026')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 1. PURPOSE & OVERVIEW
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Purpose & Overview', level=1)
doc.add_paragraph(
    'CATCAT is a browser-based tool for annotating satellite imagery of ships. '
    'Analysts draw or accept bounding boxes around detected vessels, which are '
    'submitted to a shared backend and used to continuously improve an AI detection '
    'model through iterative retraining.'
)
doc.add_paragraph(
    'This SOP covers the full workflow: loading imagery, reviewing model detections, '
    'correcting annotations, and monitoring team training progress via the dashboard.'
)

doc.add_heading('Key Concepts', level=2)
bullet('Model detections — Bounding boxes drawn automatically by the AI model (shown in green).')
bullet('User corrections — Boxes drawn manually by the analyst when the model is wrong or missed the ship (shown in yellow).')
bullet('Retraining — Every 5 accepted annotations triggers an automatic model retraining cycle.')
bullet('mAP50 — The model accuracy metric displayed on the dashboard. Higher is better (0–100%).')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM REQUIREMENTS
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('2. System Requirements', level=1)
bullet('Browser: Google Chrome or Microsoft Edge (Chromium). Firefox and Safari have limited file-save support.')
bullet('Internet connection required for model submission and dashboard.')
bullet('No software installation required — runs entirely in the browser.')
bullet('Recommended screen resolution: 1920×1080 or higher.')

note('Automatic local saving of processed crops requires Chrome or Edge. '
     'Other browsers will download files to your default Downloads folder instead.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 3. ACCESSING THE TOOL
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Accessing the Tool', level=1)
doc.add_paragraph('Open a Chromium-based browser and navigate to:')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('https://zzzstolenzzz.github.io/CATCAT/')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)

doc.add_paragraph()
doc.add_paragraph('The tool loads in Modern theme by default (dark blue interface).')

placeholder('CATCAT homepage — full interface on first load, Modern theme')

# ════════════════════════════════════════════════════════════════════════════
# 4. INTERFACE OVERVIEW
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Interface Overview', level=1)

placeholder('Annotated interface diagram with labeled sections')

doc.add_heading('Header Bar', level=2)
bullet('CATCAT title — top left.')
bullet('Session count — number of images accepted in the current browser session.')
bullet('mAP50 — current model accuracy percentage.')
bullet('Model status — "Model ready" (green) or "Loading model…"')
bullet('Theme toggle — switch between Stealth (dark) and Modern (dark blue) themes.')
bullet('Dashboard link — opens the team training dashboard.')

doc.add_heading('Toolbar', level=2)
bullet('Load Images — opens a folder picker to select your image folder.')
bullet('Accept [Enter] — accepts the current image annotation and advances.')
bullet('Clear — removes all bounding boxes from the current image.')
bullet('Back — returns to the previous image.')
bullet('Skip — skips the current image without submitting an annotation.')
bullet('Image enhancement controls (Bright, Contrast, Sharpen, A.Contrast, A.Color, Reset).')

doc.add_heading('Canvas Area', level=2)
bullet('Displays the current satellite image.')
bullet('Inch rulers appear on the top and left edges of the image.')
bullet('Green box = model detection. Yellow box = user-drawn box.')
bullet('Dotted box = 1-inch expansion around the bounding box (defines the saved crop region).')
bullet('Crosshair cursor assists with precise box placement.')

doc.add_heading('Status Bar', level=2)
bullet('Left: status messages (loading, saving, submission result).')
bullet('Center: image counter (e.g., "3 / 47").')
bullet('Right: version number and "sample images" download link.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 5. LOADING IMAGES
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Loading Images', level=1)

doc.add_heading('5.1  Using Your Own Images', level=2)
numbered('Click Load Images in the toolbar.')
numbered('A folder picker dialog appears. Navigate to the folder containing your satellite images.')
numbered('Select the folder and click "Select Folder".')
numbered('CATCAT loads all images (JPEG, PNG, WEBP, TIFF, BMP) from the folder and begins with the first one.')

placeholder('Folder picker dialog — selecting an image folder')

note('The folder you select is also where processed crops will be saved automatically '
     '(in a "Processed" subfolder). You do not need to select a save location separately.')

doc.add_heading('5.2  Using Sample Images (Demo)', level=2)
numbered('Click the small "sample images" link in the bottom-right of the status bar.')
numbered('This downloads a ZIP file containing 30 sample ship images across 5 vessels.')
numbered('Extract the ZIP to a folder on your computer.')
numbered('Click Load Images and select that extracted folder.')

placeholder('"sample images" link location — bottom-right of status bar')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 6. ANNOTATING IMAGES
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('6. Annotating Images', level=1)
doc.add_paragraph(
    'After loading, the tool automatically runs the AI model on each image and draws '
    'green bounding boxes around detected ships. The analyst reviews each detection '
    'and either accepts it or corrects it.'
)

doc.add_heading('6.1  Reviewing a Model Detection', level=2)
doc.add_paragraph('When the model detects a ship, a green bounding box appears over the vessel.')

placeholder('Image with green model detection box and dotted expansion box')

numbered('Examine the bounding box. Verify it correctly surrounds the ship.')
numbered('If the box is correct, press Enter or click Accept to submit the annotation.')
numbered('The image is submitted to the backend and the tool advances to the next image.')

tip('Check that the dotted box (1-inch expansion) captures enough context around the ship. '
    'This dotted region is what gets saved as the processed crop.')

doc.add_heading('6.2  Correcting a Wrong or Missed Detection', level=2)
doc.add_paragraph(
    'If the model box is inaccurate, or if the model missed the ship entirely, '
    'draw your own box:'
)
numbered('Click Clear to remove the existing green box (if any).')
numbered('Click and drag on the canvas to draw a new yellow bounding box around the ship.')
numbered('Adjust if needed — drawing a new box replaces the previous one.')
numbered('Press Enter or click Accept to submit your correction.')

placeholder('User drawing a yellow correction box over a ship')

note('User-drawn corrections are flagged in the system as "User ✓" and are especially '
     'valuable for model training — they teach the model where it went wrong.')

doc.add_heading('6.3  Skipping an Image', level=2)
doc.add_paragraph(
    'If an image is ambiguous, cloud-obscured, or does not contain a ship, click Skip. '
    'The image is not submitted and the tool advances. You can also use Back to revisit a previous image.'
)

placeholder('Skip and Back buttons highlighted in toolbar')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 7. IMAGE ENHANCEMENT CONTROLS
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Image Enhancement Controls', level=1)
doc.add_paragraph(
    'The toolbar provides real-time image enhancement controls to improve visibility '
    'of ships before annotating. Enhancements are also baked into saved crop files.'
)

placeholder('Enhancement controls — sliders and toggle checkboxes highlighted')

doc.add_heading('Brightness', level=2)
doc.add_paragraph('Adjusts overall image brightness. Default: 100. Range: 50–200.')

doc.add_heading('Contrast', level=2)
doc.add_paragraph('Adjusts overall image contrast. Default: 100. Range: 50–300.')

doc.add_heading('Sharpen', level=2)
doc.add_paragraph('Applies a sharpening filter to bring out edge detail. Default: 0 (off).')

doc.add_heading('A.Contrast (Auto Contrast)', level=2)
doc.add_paragraph(
    'Automatically stretches the luminance range of the image — darkens the darkest '
    'areas and brightens the brightest — improving visibility in low-contrast scenes. '
    'Equivalent to Photoshop Auto Contrast.'
)

placeholder('Same image with A.Contrast OFF vs ON — side by side comparison')

doc.add_heading('A.Color (Auto Color)', level=2)
doc.add_paragraph(
    'Independently stretches each color channel (Red, Green, Blue) to its full range. '
    'Useful for removing color casts and revealing detail in tinted imagery. '
    'Equivalent to Photoshop Auto Color.'
)

placeholder('Same image with A.Color OFF vs ON — side by side comparison')

doc.add_heading('Reset', level=2)
doc.add_paragraph('Resets Brightness, Contrast, and Sharpen sliders to defaults. Does not reset A.Contrast or A.Color.')

tip('For most satellite ship images, try A.Color first. It tends to produce the most '
    'visible improvement. A.Contrast is most useful on low-contrast or hazy imagery.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 8. LOCAL FILE SAVING
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('8. Local File Saving', level=1)
doc.add_paragraph(
    'When an image is accepted, CATCAT automatically saves a cropped JPEG of the '
    'ship region to a "Processed" subfolder inside your image folder.'
)

doc.add_heading('8.1  What Gets Saved', level=2)
bullet('The saved file is the region defined by the dotted expansion box (1 inch around the bounding box, clamped to image edges).')
bullet('Filename: original filename with "_processed" appended (e.g., "ship_001_processed.jpg").')
bullet('Maximum file size: 2.5 MB (JPEG quality is stepped down automatically if needed).')
bullet('All active image enhancements (Brightness, Contrast, A.Contrast, A.Color) are baked into the saved file.')

doc.add_heading('8.2  Save Behavior by Browser', level=2)
bullet('Chrome / Edge with folder loaded via Load Images: saves automatically to Processed/ subfolder — no prompt.')
bullet('Chrome / Edge with files loaded individually: prompts once to choose an output folder, then saves automatically.')
bullet('Firefox / Safari: triggers a browser download for each accepted image.')

placeholder('Processed folder contents — showing _processed files alongside originals')

note('If you process the same folder multiple times, existing "_processed" files are overwritten silently.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 9. COMPLETING A SESSION
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('9. Completing a Session', level=1)
doc.add_paragraph(
    'Continue accepting or skipping images until the status bar displays the '
    'bold green message "All done!" — indicating all images in the folder have been reviewed.'
)

placeholder('"All done!" message in bold green in the status bar')

doc.add_paragraph(
    'After completing a folder, you can load a new folder of images by clicking '
    'Load Images again at any time.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 10. DASHBOARD
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('10. Team Dashboard', level=1)
doc.add_paragraph(
    'The Dashboard provides visibility into team annotation activity, model training '
    'status, and historical performance. Access it via the "Dashboard →" link in the '
    'top-right of the main interface.'
)

placeholder('Full dashboard view — all panels visible')

doc.add_heading('10.1  Summary Stats', level=2)
bullet('Total Annotations — cumulative annotations submitted by the team.')
bullet('Queue — images awaiting the next training run.')
bullet('Training Runs — number of completed retraining cycles.')
bullet('mAP50 — current model accuracy.')

doc.add_heading('10.2  Training Engine Panel', level=2)
doc.add_paragraph(
    'Shows whether the model is actively training (ring animates) or idle. '
    'Displays current epoch, loss, and time elapsed during training. '
    'Also shows how many more annotations are needed before the next training run triggers.'
)

placeholder('Training engine panel — TRAINING state with animated ring')
placeholder('Training engine panel — IDLE state')

doc.add_heading('10.3  mAP50 Chart', level=2)
doc.add_paragraph(
    'Line chart showing model accuracy (mAP50) over time across all training runs. '
    'An upward trend indicates the model is improving with more annotations.'
)

placeholder('mAP50 line chart — showing upward trend across multiple runs')

doc.add_heading('10.4  Training Runs Table', level=2)
doc.add_paragraph('Shows one row per completed training run, including:')
bullet('Run number and date/time.')
bullet('mAP50 score for that run.')
bullet('Annotations — number of new annotations in that run (not cumulative).')
bullet('Total training time and seconds-per-epoch.')
bullet('Click any row to expand per-image details.')

placeholder('Training runs table — expanded row showing per-image details')

doc.add_heading('Per-Image Details', level=3)
doc.add_paragraph('Expanding a run row shows each image in that training batch:')
bullet('Image filename.')
bullet('Detections — number of ships the model found.')
bullet('Confidence — model confidence score (higher = more certain).')
bullet('Label — "Model ✓" if the model detection was accepted; "User ✓" if the analyst drew their own box.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 11. STANDARD WORKFLOW SUMMARY
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('11. Standard Workflow Summary', level=1)

numbered('Open Chrome or Edge and navigate to the CATCAT URL.')
numbered('Click Load Images and select your image folder.')
numbered('(Optional) Enable A.Color and/or A.Contrast for better visibility.')
numbered('Review the model detection (green box).')
numbered('If correct — press Enter to accept.')
numbered('If wrong — click Clear, draw your own yellow box, then press Enter.')
numbered('If no ship / ambiguous — click Skip.')
numbered('Repeat steps 4–7 until "All done!" appears.')
numbered('Check the Dashboard to confirm submissions and monitor model accuracy.')

placeholder('Workflow flowchart: Load → Review → Accept / Correct / Skip → All done → Dashboard')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 12. TROUBLESHOOTING
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('12. Troubleshooting', level=1)

doc.add_heading('"Model unavailable" or no green boxes appear', level=2)
bullet('The AI model failed to load. Check internet connection.')
bullet('Edge users: Tracking Prevention may block the model CDN. Try adding the site to trusted sites, or use Chrome.')
bullet('The tool still works in annotation-only mode — you can draw boxes manually.')

doc.add_heading('Files not saving to Processed/ folder', level=2)
bullet('Ensure you loaded images via the folder picker (Load Images → select folder), not by dragging individual files.')
bullet('If using Firefox/Safari, files download to your browser\'s default Downloads folder instead.')
bullet('Check that the browser was granted permission to access the folder.')

doc.add_heading('Confidence showing as low (10–20%)', level=2)
bullet('Low confidence is expected for a model trained on a small dataset.')
bullet('Confidence improves as more annotations are submitted and training runs accumulate.')
bullet('A low-confidence detection can still be correct — assess visually.')

doc.add_heading('Dashboard not updating', level=2)
bullet('The dashboard refreshes automatically every 10 seconds (stats) and 60 seconds (history).')
bullet('Click the browser refresh button to force an immediate reload.')
bullet('If training is active, wait for it to complete before the new run appears in the table.')

doc.add_heading('"All done!" appeared but I missed an image', level=2)
bullet('Click Load Images and reload the same folder. All images will be re-queued.')
bullet('Previously accepted images can be re-accepted — submissions stack cumulatively.')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 13. QUICK REFERENCE
# ════════════════════════════════════════════════════════════════════════════
doc.add_heading('13. Quick Reference', level=1)

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Action'
hdr[1].text = 'How'
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True

rows = [
    ('Load images', 'Click Load Images → select folder'),
    ('Accept annotation', 'Press Enter  or  click Accept'),
    ('Draw correction box', 'Click and drag on the canvas'),
    ('Clear all boxes', 'Click Clear'),
    ('Go back one image', 'Click Back'),
    ('Skip image', 'Click Skip'),
    ('Toggle Auto Contrast', 'Check/uncheck A.Contrast in toolbar'),
    ('Toggle Auto Color', 'Check/uncheck A.Color in toolbar'),
    ('Reset enhancements', 'Click Reset'),
    ('Open dashboard', 'Click "Dashboard →" (top right)'),
    ('Download sample images', 'Click "sample images" (bottom right)'),
]
for action, how in rows:
    row = table.add_row().cells
    row[0].text = action
    row[1].text = how

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# Save
# ════════════════════════════════════════════════════════════════════════════
out = r'C:\Repos\CATCAT\CATCAT_SOP.docx'
doc.save(out)
print(f'Saved: {out}')
